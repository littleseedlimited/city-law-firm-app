"""
City Law Firm - Main Telegram Bot
Handles all bot operations, commands, and automation
"""
import os
import logging
from datetime import datetime, timedelta
import sys
import os
# Fix import path to allow importing from database directory
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, WebAppInfo
from telegram.ext import (
    Application, CommandHandler, MessageHandler, CallbackQueryHandler,
    ConversationHandler, filters, ContextTypes, PicklePersistence
)
from dotenv import load_dotenv
import asyncio
import pytz
from dotenv import load_dotenv

# Load environment variables (MUST BE FIRST)
load_dotenv('config/.env')

from database.models import (
    init_db, get_session, User, Case, CourtDate, 
    TimeEntry, LeaveRequest, Notification, ComplianceTask, Document, PaymentRequest
)
import uuid
from bot.scheduler import start_scheduler




# Configure logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Initialize database
engine = init_db()

# Onboarding conversation states
(ONBOARD_NAME, ONBOARD_EMAIL, ONBOARD_PHONE, ONBOARD_DEPARTMENT,
 ONBOARD_POSITION, ONBOARD_SPECIALIZATION, ONBOARD_BAR_NUMBER,
 ONBOARD_ADDRESS, ONBOARD_EMERGENCY_NAME, ONBOARD_EMERGENCY_PHONE,
 ONBOARD_PHOTO, ONBOARD_LOCATION, ONBOARD_CONFIRM) = range(13)

# Profile editing conversation states
(EDIT_PHONE, EDIT_EMAIL, EDIT_ADDRESS) = range(13, 16)

# Department configuration
DEPARTMENTS = {
    'partners': {'name': 'Partners & Management', 'icon': 'ðŸ‘”', 'max_members': 3},
    'litigation': {'name': 'Litigation Department', 'icon': 'âš–ï¸', 'max_members': 8},
    'corporate': {'name': 'Corporate Law', 'icon': 'ðŸ¢', 'max_members': 6},
    'family': {'name': 'Family Law', 'icon': 'ðŸ‘¨â€ðŸ‘©â€ðŸ‘§', 'max_members': 4},
    'criminal': {'name': 'Criminal Defense', 'icon': 'ðŸ”’', 'max_members': 5},
    'admin': {'name': 'Administration & HR', 'icon': 'ðŸ“‹', 'max_members': 4}
}


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start command - Welcome message"""
    user = update.effective_user
    session = get_session(engine)
    
    # Check if user exists
    db_user = session.query(User).filter_by(telegram_id=user.id).first()
    
    if db_user and db_user.onboarding_completed:
        # Fetch latest broadcast
        latest_broadcast = session.query(Notification).filter_by(notification_type='broadcast').order_by(Notification.created_at.desc()).first()
        broadcast_param = ""
        if latest_broadcast:
            import urllib.parse
            msg = urllib.parse.quote(latest_broadcast.message)
            time = urllib.parse.quote(latest_broadcast.created_at.strftime('%d-%m-%Y %H:%M'))
            broadcast_param = f"?broadcast={msg}&time={time}"

        # Existing user - show main menu
        keyboard = [
            [InlineKeyboardButton("ðŸ“± Open Virtual Office", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL') + broadcast_param))],
            [InlineKeyboardButton("ðŸ“‹ My Agenda", callback_data='my_agenda')],
            [InlineKeyboardButton("ðŸ“Š Dashboard", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL') + broadcast_param))],
            [InlineKeyboardButton("â„¹ï¸ Help", callback_data='help')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            f"ðŸ‘‹ Welcome back, {db_user.full_name}!\n\n"
            f"ðŸ¢ **City Law Firm Virtual Office**\n"
            f"Departments: {db_user.departments}\n"
            f"Position: {db_user.position}\n\n"
            f"What would you like to do today?",
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
    else:
        # New user - start onboarding
        keyboard = [
            [InlineKeyboardButton("âœ… Start Onboarding", callback_data='start_onboarding')],
            [InlineKeyboardButton("â„¹ï¸ Learn More", callback_data='learn_more')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            f"ðŸ‘‹ Welcome to **City Law Firm Virtual Office**!\n\n"
            f"I'm your virtual assistant, designed to help you navigate "
            f"our digital workspace efficiently.\n\n"
            f"ðŸ¢ **What I can help you with:**\n"
            f"â€¢ Onboarding new staff members\n"
            f"â€¢ Managing cases and court dates\n"
            f"â€¢ Tracking billable time\n"
            f"â€¢ Handling leave requests\n"
            f"â€¢ Notifications and reminders\n"
            f"â€¢ Accessing resources and documents\n\n"
            f"Let's get you set up!",
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
    
    session.close()


async def start_onboarding_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start the onboarding process"""
    query = update.callback_query
    await query.answer()
    
    await query.edit_message_text(
        "ðŸŽ¯ **Welcome to City Law Firm Onboarding**\n\n"
        "I'll guide you through a quick setup process to get you started. "
        "This will take about 5 minutes.\n\n"
        "You can update any information later using the /profile command.\n\n"
        "Let's begin! ðŸ‘‡\n\n"
        "**Step 1/10:** Please enter your full name:",
        parse_mode='Markdown'
    )
    
    return ONBOARD_NAME


async def onboard_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect full name"""
    context.user_data['full_name'] = update.message.text
    
    await update.message.reply_text(
        f"Great, {update.message.text}! âœ…\n\n"
        f"**Step 2/10:** Please enter your work email address:",
        parse_mode='Markdown'
    )
    
    return ONBOARD_EMAIL


async def onboard_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect email"""
    email = update.message.text
    
    # Basic email validation
    if '@' not in email or '.' not in email:
        await update.message.reply_text(
            "âš ï¸ Please enter a valid email address (e.g., john.doe@citylawfirm.com):"
        )
        return ONBOARD_EMAIL
    
    context.user_data['email'] = email
    
    await update.message.reply_text(
        f"Email saved! âœ…\n\n"
        f"**Step 3/10:** Please enter your phone number:",
        parse_mode='Markdown'
    )
    
    return ONBOARD_PHONE


def _get_department_keyboard(selected_depts):
    """Helper to generate department keyboard with checkmarks"""
    keyboard = []
    for key, data in DEPARTMENTS.items():
        icon = "âœ… " if data['name'] in selected_depts else data['icon'] + " "
        keyboard.append([InlineKeyboardButton(f"{icon}{data['name']}", callback_data=f"dept_{key}")])
    
    keyboard.append([InlineKeyboardButton("âœ… Done / Continue", callback_data="dept_done")])
    return keyboard


async def onboard_phone(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect phone number"""
    context.user_data['phone'] = update.message.text
    context.user_data['departments'] = []  # Initialize empty list
    
    # Department selection keyboard
    keyboard = _get_department_keyboard([])
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        f"Phone number saved! âœ…\n\n"
        f"**Step 4/10:** Select your department(s) (Select all that apply):\n"
        f"Click 'Done' when finished.",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    return ONBOARD_DEPARTMENT


async def onboard_department(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle multi-department selection"""
    query = update.callback_query
    await query.answer()
    
    data = query.data
    
    if data == 'dept_done':
        if not context.user_data.get('departments'):
            await query.answer("Please select at least one department!", show_alert=True)
            return ONBOARD_DEPARTMENT
            
        selected_str = ", ".join(context.user_data['departments'])
        await query.edit_message_text(
            f"Departments: {selected_str} âœ…\n\n"
            f"**Step 5/10:** Please enter your position/title:\n"
            f"(e.g., Associate Attorney, Partner, Paralegal, HR Manager)",
            parse_mode='Markdown'
        )
        return ONBOARD_POSITION
        
    # Toggle selection
    dept_key = data.replace('dept_', '')
    dept_name = DEPARTMENTS[dept_key]['name']
    
    current_depts = context.user_data.get('departments', [])
    if dept_name in current_depts:
        current_depts.remove(dept_name)
    else:
        current_depts.append(dept_name)
    
    context.user_data['departments'] = current_depts
    
    # Update keyboard
    reply_markup = InlineKeyboardMarkup(_get_department_keyboard(current_depts))
    await query.edit_message_reply_markup(reply_markup=reply_markup)
    
    return ONBOARD_DEPARTMENT


async def onboard_position(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect position"""
    context.user_data['position'] = update.message.text
    
    await update.message.reply_text(
        f"Position saved! âœ…\n\n"
        f"**Step 6/10:** What's your area of specialization?\n"
        f"(e.g., Contract Law, Criminal Defense, Family Law, or type 'N/A' if not applicable)",
        parse_mode='Markdown'
    )
    
    return ONBOARD_SPECIALIZATION


async def onboard_specialization(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect specialization"""
    context.user_data['specialization'] = update.message.text
    
    await update.message.reply_text(
        f"Specialization noted! âœ…\n\n"
        f"**Step 7/10:** Please enter your Bar Number:\n"
        f"(or type 'N/A' if not applicable)",
        parse_mode='Markdown'
    )
    
    return ONBOARD_BAR_NUMBER


async def onboard_bar_number(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect bar number"""
    bar_number = update.message.text
    context.user_data['bar_number'] = None if bar_number.upper() == 'N/A' else bar_number
    
    await update.message.reply_text(
        f"Bar number saved! âœ…\n\n"
        f"**Step 8/10:** Please enter your current address:",
        parse_mode='Markdown'
    )
    
    return ONBOARD_ADDRESS


async def onboard_address(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect address"""
    context.user_data['address'] = update.message.text
    
    await update.message.reply_text(
        f"Address saved! âœ…\n\n"
        f"**Step 9/10:** Emergency Contact Name:",
        parse_mode='Markdown'
    )
    
    return ONBOARD_EMERGENCY_NAME


async def onboard_emergency_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect emergency contact name"""
    context.user_data['emergency_contact_name'] = update.message.text
    
    await update.message.reply_text(
        f"Contact name saved! âœ…\n\n"
        f"**Step 10/10:** Emergency Contact Phone:",
        parse_mode='Markdown'
    )
    
    return ONBOARD_EMERGENCY_PHONE


async def onboard_emergency_phone(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect emergency contact phone and request photo"""
    context.user_data['emergency_contact_phone'] = update.message.text
    
    await update.message.reply_text(
        "ðŸ“¸ **Almost done! Please upload a profile photo.**\n\n"
        "This will be used for your staff ID and the virtual office dashboard.\n"
        "_Please send a photo (not a file)._",
        parse_mode='Markdown'
    )
    
    return ONBOARD_PHOTO


async def onboard_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect profile photo and request location"""
    photo_file = await update.message.photo[-1].get_file()
    context.user_data['photo_file_id'] = photo_file.file_id
    
    # Request location
    keyboard = [[
        InlineKeyboardButton("ðŸ“ Share Location", callback_data='skip_location')  # Placeholder, actual location request needs ReplyKeyboardMarkup
    ]]
    # Actually, for location we need ReplyKeyboardMarkup
    from telegram import KeyboardButton, ReplyKeyboardMarkup
    
    location_keyboard = ReplyKeyboardMarkup(
        [[KeyboardButton("ðŸ“ Share Current Location", request_location=True)]],
        one_time_keyboard=True,
        resize_keyboard=True
    )
    
    await update.message.reply_text(
        "ðŸ“ **Last Step: Location Access**\n\n"
        "To enable the **Staff Status Dashboard** and geotagging features, "
        "please share your current location.\n\n"
        "This helps us coordinate field work and office presence.",
        reply_markup=location_keyboard,
        parse_mode='Markdown'
    )
    
    return ONBOARD_LOCATION


async def onboard_location(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect location and confirm details"""
    if update.message.location:
        context.user_data['latitude'] = update.message.location.latitude
        context.user_data['longitude'] = update.message.location.longitude
    
    data = context.user_data
    depts_str = ", ".join(data['departments'])
    
    keyboard = [[InlineKeyboardButton("âœ… Confirm & Submit", callback_data='confirm_onboarding')]]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    # Remove ReplyKeyboard
    from telegram import ReplyKeyboardRemove
    await update.message.reply_text("âœ… Location received!", reply_markup=ReplyKeyboardRemove())
    
    await update.message.reply_text(
        f"ðŸŽ‰ **Please confirm your details:**\n\n"
        f"ðŸ‘¤ **Name:** {data['full_name']}\n"
        f"ðŸ“§ **Email:** {data['email']}\n"
        f"ðŸ“± **Phone:** {data['phone']}\n"
        f"ðŸ¢ **Departments:** {depts_str}\n"
        f"ðŸ’¼ **Position:** {data['position']}\n"
        f"âš–ï¸ **Specialization:** {data['specialization']}\n"
        f"ðŸ†” **Bar Number:** {data['bar_number'] or 'N/A'}\n"
        f"ðŸ  **Address:** {data['address']}\n"
        f"ðŸš‘ **Emergency:** {data['emergency_contact_name']} ({data['emergency_contact_phone']})\n"
        f"ðŸ“¸ **Photo:** Received\n"
        f"ðŸ“ **Location:** {data.get('latitude', 'N/A')}, {data.get('longitude', 'N/A')}",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    return ONBOARD_CONFIRM


async def confirm_onboarding(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Save user data and complete onboarding"""
    query = update.callback_query
    await query.answer()
    
    user = update.effective_user
    session = get_session(engine)
    
    try:
        # Check if user exists
        db_user = session.query(User).filter_by(telegram_id=user.id).first()
        
        if not db_user:
            db_user = User(telegram_id=user.id)
            session.add(db_user)
        
        # Update user information
        data = context.user_data
        db_user.username = user.username
        db_user.full_name = data['full_name']
        db_user.email = data['email']
        db_user.phone = data['phone']
        db_user.departments = ", ".join(data['departments'])  # Save as comma-separated string
        db_user.position = data['position']
        db_user.specialization = data['specialization']
        db_user.bar_number = data.get('bar_number')
        db_user.address = data['address']
        db_user.emergency_contact_name = data['emergency_contact_name']
        db_user.emergency_contact_phone = data['emergency_contact_phone']
        db_user.onboarding_completed = True
        db_user.onboarding_completed_at = datetime.utcnow()
        db_user.role = 'staff'
        db_user.status = 'active'
        
        session.commit()
        
        # Welcome message with main menu
        keyboard = [
            [InlineKeyboardButton("ðŸ“± Open Virtual Office", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL')))],
            [InlineKeyboardButton("ðŸ“š Quick Start Guide", callback_data='quick_start')],
            [InlineKeyboardButton("â„¹ï¸ Help & Commands", callback_data='help')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await query.edit_message_text(
            f"ðŸŽ‰ **Welcome to City Law Firm, {data['full_name']}!**\n\n"
            f"Your onboarding is complete! You now have access to:\n\n"
            f"âœ… Virtual Office Dashboard\n"
            f"âœ… {', '.join(data['departments'])} Channels\n"
            f"âœ… Case Management Tools\n"
            f"âœ… Time Tracking System\n"
            f"âœ… Document Repository\n"
            f"âœ… All Department Bots\n\n"
            f"Use /help to see all available commands.\n\n"
            f"Let's get started! ðŸ‘‡",
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        
        # Clear conversation data
        context.user_data.clear()
        
        return ConversationHandler.END
        
    except Exception as e:
        logger.error(f"Error saving user data: {e}")
        await query.edit_message_text(
            "âŒ Sorry, there was an error saving your information. "
            "Please try again later or contact an administrator."
        )
        session.rollback()
        return ConversationHandler.END
    finally:
        session.close()


async def dashboard_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Show user dashboard with real statistics"""
    query = update.callback_query
    await query.answer()
    
    user = update.effective_user
    session = get_session(engine)
    
    try:
        db_user = session.query(User).filter_by(telegram_id=user.id).first()
        
        if not db_user or not db_user.onboarding_completed:
            await query.edit_message_text(
                "âš ï¸ Please complete onboarding first using /start"
            )
            return
        
        # Get statistics
        from sqlalchemy import func, and_
        from datetime import datetime, timedelta
        
        today = datetime.now().date()
        week_from_now = today + timedelta(days=7)
        
        # Count active cases
        active_cases = session.query(Case).filter(
            Case.assigned_to == db_user.id,
            Case.status.in_(['open', 'in_progress'])
        ).count()
        
        # Count pending time entries (today)
        time_entries_today = session.query(TimeEntry).filter(
            TimeEntry.user_id == db_user.id,
            func.date(TimeEntry.date) == today
        ).count()
        
        # Get upcoming court dates
        upcoming_court_dates = session.query(CourtDate).join(Case).filter(
            Case.assigned_to == db_user.id,
            CourtDate.hearing_date >= datetime.now(),
            CourtDate.hearing_date <= datetime.combine(week_from_now, datetime.max.time())
        ).count()
        
        # Get leave request status
        pending_leave = session.query(LeaveRequest).filter(
            LeaveRequest.user_id == db_user.id,
            LeaveRequest.status == 'pending'
        ).count()
        
        # Get total billable hours this month
        from datetime import date
        first_day_month = date.today().replace(day=1)
        total_hours = session.query(func.sum(TimeEntry.duration)).filter(
            TimeEntry.user_id == db_user.id,
            TimeEntry.date >= first_day_month,
            TimeEntry.billable == True
        ).scalar() or 0
        
        # Build dashboard message
        dashboard_text = (
            f"ðŸ“Š **Dashboard - {db_user.full_name}**\n\n"
            f"ðŸ¢ **Department(s):** {db_user.departments}\n"
            f"ðŸ’¼ **Position:** {db_user.position}\n\n"
            f"**ðŸ“ˆ Your Statistics:**\n"
            f"â€¢ Active Cases: {active_cases}\n"
            f"â€¢ Time Entries Today: {time_entries_today}\n"
            f"â€¢ Upcoming Court Dates (7 days): {upcoming_court_dates}\n"
            f"â€¢ Pending Leave Requests: {pending_leave}\n"
            f"â€¢ Billable Hours (This Month): {total_hours:.1f}h\n\n"
            f"_Last updated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}_"
        )
        
        keyboard = [
            [InlineKeyboardButton("ðŸ“± Open Virtual Office", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL')))],
            [InlineKeyboardButton("ðŸ“‹ My Agenda", callback_data='my_agenda')],
            [InlineKeyboardButton("ðŸ”„ Refresh", callback_data='dashboard')],
            [InlineKeyboardButton("Â« Back", callback_data='back_to_start')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await query.edit_message_text(
            dashboard_text,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        
    except Exception as e:
        logger.error(f"Error loading dashboard: {e}")
        await query.edit_message_text(
            "âŒ Error loading dashboard. Please try again later."
        )
    finally:
        session.close()





async def newcase(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/newcase - Register a new case"""
    keyboard = [
        [InlineKeyboardButton("ðŸ“± Open Case Form", web_app=WebAppInfo(url=f"{os.getenv('MINI_APP_URL')}?view=newcase"))],
        [InlineKeyboardButton("âŒ Cancel", callback_data='cancel')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        "ðŸ“ **Register New Case**\n\n"
        "Click below to open the case registration form in the Virtual Office app.\n\n"
        "You can also use the quick format:\n"
        "`/newcase [case_number] [client_name] [case_type]`",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def casestatus(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/casestatus [id] - Check status of cases"""
    session = get_session(engine)
    
    try:
        # If no args, show all cases
        if not context.args:
            cases = session.query(Case).order_by(Case.updated_at.desc()).all()
            
            if not cases:
                await update.message.reply_text(
                    "ðŸ“‹ **No cases found**\n\n"
                    "There are currently no cases in the system."
                )
                return
            
            # Build message with all cases
            message = f"ðŸ“‹ **All Cases ({len(cases)} total)**\n\n"
            
            keyboard = []
            for case in cases[:10]:  # Limit to 10 most recent
                status_emoji = "ðŸŸ¢" if case.status == "active" else "ðŸŸ¡" if case.status == "pending" else "ðŸ”´"
                message += (
                    f"{status_emoji} **{case.case_number}**\n"
                    f"   {case.title}\n"
                    f"   Client: {case.client_name}\n"
                    f"   Status: {case.status.upper()}\n"
                    f"   Updated: {case.updated_at.strftime('%b %d, %Y')}\n\n"
                )
                
                # Add button for each case
                keyboard.append([InlineKeyboardButton(
                    f"ðŸ“„ View {case.case_number}", 
                    callback_data=f'view_case_{case.id}'
                )])
            
            if len(cases) > 10:
                message += f"\n_Showing 10 of {len(cases)} cases_"
            
            reply_markup = InlineKeyboardMarkup(keyboard)
            await update.message.reply_text(
                message,
                reply_markup=reply_markup,
                parse_mode='Markdown'
            )
            return
        
        # If args provided, show specific case
        case_number = context.args[0]
        case = session.query(Case).filter_by(case_number=case_number).first()
        
        if not case:
            await update.message.reply_text(
                f"âŒ Case '{case_number}' not found.\n\n"
                f"Please check the case number and try again."
            )
            return
        
        # Build case status message
        message = (
            f"ðŸ“‹ **Case Status Report**\n\n"
            f"**Case Number:** {case.case_number}\n"
            f"**Title:** {case.title}\n"
            f"**Client:** {case.client_name}\n"
            f"**Type:** {case.case_type}\n"
            f"**Status:** {case.status.upper()}\n"
            f"**Priority:** {case.priority.upper()}\n"
            f"**Department:** {case.department}\n\n"
        )
        
        if case.assigned_to_user:
            message += f"**Assigned To:** {case.assigned_to_user.full_name}\n\n"
        
        if case.next_court_date:
            message += f"ðŸ“… **Next Court Date:** {case.next_court_date.strftime('%B %d, %Y at %I:%M %p')}\n\n"
        
        if case.deadline:
            message += f"â° **Deadline:** {case.deadline.strftime('%B %d, %Y')}\n\n"
        
        if case.description:
            message += f"**Description:**\n{case.description[:200]}...\n\n"
        
        message += f"_Last updated: {case.updated_at.strftime('%B %d, %Y')}_"
        
        keyboard = [
            [InlineKeyboardButton("ðŸ“ Add Time Entry", callback_data=f'time_entry_{case.id}')],
            [InlineKeyboardButton("ðŸ“„ View Documents", callback_data=f'docs_{case.id}')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            message,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        
    finally:
        session.close()


async def requestleave(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/requestleave - Submit leave request"""
    keyboard = [
        [InlineKeyboardButton("ðŸ“± Open Leave Form", web_app=WebAppInfo(url=f"{os.getenv('MINI_APP_URL')}/leave-request"))],
        [InlineKeyboardButton("âŒ Cancel", callback_data='cancel')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        "ðŸ–ï¸ **Request Time Off**\n\n"
        "Click below to open the leave request form.\n\n"
        "Leave types available:\n"
        "â€¢ Vacation\n"
        "â€¢ Sick Leave\n"
        "â€¢ Personal Leave\n"
        "â€¢ Emergency Leave",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def resources(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/resources [topic] - Search legal resources"""
    if not context.args:
        await update.message.reply_text(
            "ðŸ“š **Legal Resources**\n\n"
            "Usage: `/resources [search_topic]`\n"
            "Example: `/resources contract templates`\n\n"
            "Available resources:\n"
            "â€¢ Legal templates\n"
            "â€¢ Case law databases\n"
            "â€¢ Statute references\n"
            "â€¢ Internal policies\n"
            "â€¢ Training materials",
            parse_mode='Markdown'
        )
        return
    
    search_term = ' '.join(context.args)
    
    keyboard = [
        [InlineKeyboardButton("ðŸ” Search Resources", web_app=WebAppInfo(url=f"{os.getenv('MINI_APP_URL')}/resources?q={search_term}"))],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        f"ðŸ” Searching for: **{search_term}**\n\n"
        f"Click below to view search results:",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def emergency(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/emergency - Alert partners about urgent matter"""
    user = update.effective_user
    session = get_session(engine)
    
    try:
        db_user = session.query(User).filter_by(telegram_id=user.id).first()
        
        if not db_user:
            await update.message.reply_text("âš ï¸ Please complete onboarding first.")
            return
        
        # Get all partners
        partners = session.query(User).filter_by(role='partner', status='active').all()
        
        if not partners:
            await update.message.reply_text(
                "âš ï¸ No partners available to notify."
            )
            return
        
        keyboard = [
            [InlineKeyboardButton("âœ… Confirm Emergency Alert", callback_data='confirm_emergency')],
            [InlineKeyboardButton("âŒ Cancel", callback_data='cancel')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            "ðŸš¨ **Emergency Alert**\n\n"
            f"This will immediately notify all {len(partners)} partners.\n\n"
            "Please confirm this is an urgent matter:",
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        
    finally:
        session.close()


async def refer(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/refer [case] - Refer case to another department"""
    if not context.args:
        await update.message.reply_text(
            "ðŸ”„ **Refer Case**\n\n"
            "Usage: `/refer [case_number]`\n"
            "Example: `/refer CL-2025-001`",
            parse_mode='Markdown'
        )
        return
    
    case_number = context.args[0]
    
    keyboard = []
    for key, dept in DEPARTMENTS.items():
        keyboard.append([InlineKeyboardButton(dept['name'], callback_data=f'refer_{case_number}_{key}')])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        f"ðŸ”„ **Refer Case {case_number}**\n\n"
        f"Select the department to refer this case to:",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/help - Show all available commands"""
    help_text = (
        "ðŸ¤– **City Law Firm Bot - Command Reference**\n\n"
        
        "**ðŸ“± MAIN COMMANDS**\n"
        "/start - Start bot and access main menu\n"
        "/myagenda - View your daily schedule and court dates\n"
        "/help - Show this help message\n\n"
        
        "**ðŸ“‹ CASE MANAGEMENT**\n"
        "/newcase - Register a new case\n"
        "/casestatus [id] - Check status of a case\n"
        "/refer [case] - Refer case to another department\n\n"
        
        "**ðŸ‘¤ PERSONAL**\n"
        "/profile - View/edit your profile\n"
        "/requestleave - Submit time-off request\n"
        "/mytime - View your billable hours\n\n"
        
        "**ðŸ“š RESOURCES**\n"
        "/resources [topic] - Search legal resources\n"
        "/documents - Access document repository\n\n"
        
        "**ðŸš¨ ALERTS**\n"
        "/emergency - Alert partners about urgent matter\n"
        "/notifications - View all notifications\n\n"
        
        "**â„¹ï¸ INFORMATION**\n"
        "/directory - Staff directory\n"
        "/departments - View all departments\n"
        "/about - About City Law Firm\n\n"
        
        "**ðŸ¤– SPECIALIZED BOTS**\n"
        "@CaseManagerBot - Case tracking and management\n"
        "@CourtDateBot - Court appearance reminders\n"
        "@ClientUpdateBot - Client communication tracking\n"
        "@BillingBot - Time tracking and billing\n"
        "@ComplianceBot - Compliance monitoring\n"
        "@HRBot - HR services\n"
        "@LeaveTrackerBot - Leave management\n\n"
        
        "Need more help? Contact IT support or your department head."
    )
    
    keyboard = [
        [InlineKeyboardButton("ðŸ“± Open Virtual Office", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL')))],
        [InlineKeyboardButton("ðŸ“š User Guide", callback_data='user_guide')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        help_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def profile(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/profile - View/edit profile"""
    user = update.effective_user
    session = get_session(engine)
    
    try:
        db_user = session.query(User).filter_by(telegram_id=user.id).first()
        
        if not db_user:
            await update.message.reply_text("âš ï¸ Please complete onboarding first using /start")
            return
        
        profile_text = (
            f"ðŸ‘¤ **Your Profile**\n\n"
            f"**Personal Information:**\n"
            f"â€¢ Name: {db_user.full_name}\n"
            f"â€¢ Email: {db_user.email}\n"
            f"â€¢ Phone: {db_user.phone}\n\n"
            f"**Professional:**\n"
            f"â€¢ Department: {db_user.department}\n"
            f"â€¢ Position: {db_user.position}\n"
            f"â€¢ Specialization: {db_user.specialization or 'N/A'}\n"
            f"â€¢ Bar Number: {db_user.bar_number or 'N/A'}\n\n"
            f"â€¢ Employee ID: {db_user.employee_id or 'Pending'}\n"
            f"â€¢ Status: {db_user.status.title()}\n\n"
            f"_Member since: {db_user.join_date.strftime('%B %d, %Y')}_"
        )
        
        keyboard = [
            [InlineKeyboardButton("âœï¸ Edit Profile", web_app=WebAppInfo(url=f"{os.getenv('MINI_APP_URL')}/profile"))],
            [InlineKeyboardButton("ðŸ”„ Update Information", callback_data='update_profile')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            profile_text,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        
    finally:
        session.close()


async def post_init(application: Application) -> None:
    """Initialize scheduler after application starts"""
    logger.info("ðŸ”„ Starting automated task scheduler...")
    await start_scheduler()


async def quickstart(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/quickstart - Quick Start Guide for new users"""
    guide_text = (
        "ðŸš€ **Quick Start Guide**\n\n"
        "Welcome to City Law Firm Virtual Office! Here's how to get started:\n\n"
        
        "**Step 1: Complete Onboarding**\n"
        "Send /start and click 'Start Onboarding' to set up your profile.\n\n"
        
        "**Step 2: Open Virtual Office**\n"
        "Access the Mini App for a complete dashboard with:\n"
        "â€¢ Your cases and court dates\n"
        "â€¢ Team status and locations\n"
        "â€¢ Notifications and tasks\n\n"
        
        "**Step 3: Key Commands**\n"
        "â€¢ /myagenda - View today's schedule\n"
        "â€¢ /casestatus - Check case progress\n"
        "â€¢ /profile - View/edit your profile\n"
        "â€¢ /help - See all commands\n\n"
        
        "**Step 4: Stay Connected**\n"
        "â€¢ Enable notifications for reminders\n"
        "â€¢ Share location for team visibility\n"
        "â€¢ Log billable hours regularly\n\n"
        
        "**Need Help?** Send /help or contact IT support."
    )
    
    keyboard = [
        [InlineKeyboardButton("ðŸ“± Open Virtual Office", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL')))],
        [InlineKeyboardButton("âœ… Start Onboarding", callback_data='start_onboarding')],
        [InlineKeyboardButton("ðŸ“š Full Command List", callback_data='help')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        guide_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def user_guide_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Show detailed user guide"""
    query = update.callback_query
    await query.answer()
    
    guide_text = (
        "ðŸ“š **City Law Firm - User Guide**\n\n"
        
        "**ðŸ¢ Virtual Office Mini App**\n"
        "The Mini App is your central hub for:\n"
        "â€¢ Dashboard with stats and agenda\n"
        "â€¢ Case management and tracking\n"
        "â€¢ Department information\n"
        "â€¢ Staff directory with locations\n"
        "â€¢ Profile management\n\n"
        
        "**ðŸ“‹ Case Management**\n"
        "â€¢ Register new cases via Mini App or /newcase\n"
        "â€¢ Track case status with /casestatus\n"
        "â€¢ Refer cases between departments\n\n"
        
        "**â° Time & Attendance**\n"
        "â€¢ Log billable hours through the Mini App\n"
        "â€¢ Request leave with forms in the app\n"
        "â€¢ View your agenda with /myagenda\n\n"
        
        "**ðŸ”” Notifications**\n"
        "â€¢ Court date reminders (1 day before)\n"
        "â€¢ Deadline alerts\n"
        "â€¢ System broadcasts\n\n"
        
        "**ðŸ“ Location Sharing**\n"
        "Staff can share their location for team visibility.\n"
        "View locations on the Staff Status tab.\n\n"
        
        "**ðŸ”’ Security**\n"
        "â€¢ Your data is encrypted\n"
        "â€¢ Only authorized personnel can access cases\n"
        "â€¢ All actions are logged"
    )
    
    keyboard = [
        [InlineKeyboardButton("ðŸ“± Open Virtual Office", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL')))],
        [InlineKeyboardButton("ðŸš€ Quick Start", callback_data='quickstart_cb')],
        [InlineKeyboardButton("ðŸ“ž Contact Support", url="https://t.me/LittleSeedAI")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.edit_message_text(
        guide_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def learn_more_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Show learn more info for new users"""
    query = update.callback_query
    await query.answer()
    
    info_text = (
        "â„¹ï¸ **About City Law Firm Virtual Office**\n\n"
        
        "This is a comprehensive Telegram-based office management system designed "
        "specifically for law firms.\n\n"
        
        "**Features:**\n"
        "â€¢ ðŸ“± Mini App with visual dashboard\n"
        "â€¢ ðŸ“‹ Case management and tracking\n"
        "â€¢ â° Court date reminders\n"
        "â€¢ ðŸ’¼ Billable hours logging\n"
        "â€¢ ðŸ‘¥ Staff directory with locations\n"
        "â€¢ ðŸ“¢ Notifications and broadcasts\n"
        "â€¢ ðŸ–ï¸ Leave request management\n"
        "â€¢ ðŸ“š Legal resources search\n\n"
        
        "**Departments:**\n"
        "â€¢ Partners & Management\n"
        "â€¢ Litigation\n"
        "â€¢ Corporate Law\n"
        "â€¢ Family Law\n"
        "â€¢ Criminal Defense\n"
        "â€¢ Administration & HR\n\n"
        
        "Ready to join? Click **Start Onboarding** below!"
    )
    
    keyboard = [
        [InlineKeyboardButton("âœ… Start Onboarding", callback_data='start_onboarding')],
        [InlineKeyboardButton("ðŸš€ Quick Start Guide", callback_data='quickstart_cb')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.edit_message_text(
        info_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def quickstart_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle quickstart callback for inline buttons"""
    query = update.callback_query
    await query.answer()
    
    guide_text = (
        "ðŸš€ **Quick Start Guide**\n\n"
        "Welcome to City Law Firm Virtual Office!\n\n"
        
        "**Step 1:** Complete onboarding with /start\n"
        "**Step 2:** Open the Virtual Office Mini App\n"
        "**Step 3:** Explore your dashboard\n"
        "**Step 4:** Set up notifications\n\n"
        
        "Key commands: /help, /myagenda, /profile"
    )
    
    keyboard = [
        [InlineKeyboardButton("ðŸ“± Open Virtual Office", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL')))],
        [InlineKeyboardButton("âœ… Start Onboarding", callback_data='start_onboarding')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.edit_message_text(
        guide_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )



async def myagenda_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle myagenda callback"""
    query = update.callback_query
    await query.answer()
    
    user = update.effective_user
    session = get_session(engine)
    
    try:
        db_user = session.query(User).filter_by(telegram_id=user.id).first()
        today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        
        # Get Court Dates for next 7 days
        week_from_now = today_start + timedelta(days=7)
        court_dates = session.query(CourtDate).join(Case).filter(
            Case.assigned_to == db_user.id,
            CourtDate.hearing_date >= today_start,
            CourtDate.hearing_date <= week_from_now
        ).order_by(CourtDate.hearing_date).limit(5).all()
        
        # Get Compliance Tasks
        tasks = session.query(ComplianceTask).filter(
            ComplianceTask.assigned_to == db_user.id,
            ComplianceTask.status == 'pending'
        ).order_by(ComplianceTask.deadline).limit(5).all()

        msg = f"ðŸ“‹ **My Agenda - {db_user.full_name}**\n\n"
        
        msg += "**ðŸ›ï¸ Upcoming Hearings (7 Days):**\n"
        if court_dates:
            for cd in court_dates:
                msg += f"â€¢ {cd.hearing_date.strftime('%d/%m %H:%M')} - {cd.case.case_number}\n"
        else:
            msg += "_No upcoming hearings this week._\n"
            
        msg += "\n**âœ… Pending Tasks:**\n"
        if tasks:
            for t in tasks:
                dx = t.deadline.strftime('%d/%m') if t.deadline else 'No date'
                msg += f"â€¢ {t.title} (Due {dx})\n"
        else:
            msg += "_No pending tasks._\n"

        keyboard = [
            [InlineKeyboardButton("ðŸ“± Open Full Agenda", web_app=WebAppInfo(url=f"{os.getenv('MINI_APP_URL')}"))],
            [InlineKeyboardButton("Â« Back to Menu", callback_data='help')]
        ]
        
        await query.edit_message_text(
            msg,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        
    except Exception as e:
        logger.error(f"Error in myagenda_callback: {e}")
        await query.edit_message_text("âŒ Error loading agenda.")
    finally:
        session.close()


async def help_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle help callback for inline button"""
    query = update.callback_query
    await query.answer()
    
    help_text = (
        "ðŸ¤– **Command Reference**\n\n"
        "**Main:** /start, /help, /myagenda\n"
        "**Cases:** /newcase, /casestatus\n"
        "**Personal:** /profile, /requestleave\n"
        "**Resources:** /resources\n"
        "**Alerts:** /emergency\n\n"
        "Use /help for full details."
    )
    
    keyboard = [
        [InlineKeyboardButton("ðŸ“± Open Virtual Office", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL')))],
        [InlineKeyboardButton("ðŸ“š User Guide", callback_data='user_guide')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.edit_message_text(
        help_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


def main():
    """Start the bot"""
    # Create application
    application = Application.builder().token(os.getenv('BOT_TOKEN')).post_init(post_init).build()
    
    # Callback Query Handlers (Register BEFORE ConversationHandler)
    application.add_handler(CallbackQueryHandler(dashboard_callback, pattern='^dashboard$'))
    application.add_handler(CallbackQueryHandler(doc_callback_handler, pattern='^doc_'))
    application.add_handler(CallbackQueryHandler(myagenda_callback, pattern='^my_agenda$'))
    application.add_handler(CallbackQueryHandler(help_callback, pattern='^help$'))
    
    # File Handler (Register BEFORE ConversationHandler to ensure it catches files)
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    
    # Document Follow-up Handler (Text messages - Register BEFORE ConversationHandler)
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND & ~filters.Regex('^/'), handle_document_followup))
    
    # Onboarding conversation handler
    onboarding_handler = ConversationHandler(
        entry_points=[CallbackQueryHandler(start_onboarding_callback, pattern='^start_onboarding$')],
        states={
            ONBOARD_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_name)],
            ONBOARD_EMAIL: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_email)],
            ONBOARD_PHONE: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_phone)],
            ONBOARD_DEPARTMENT: [CallbackQueryHandler(onboard_department, pattern='^dept_')],
            ONBOARD_POSITION: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_position)],
            ONBOARD_SPECIALIZATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_specialization)],
            ONBOARD_BAR_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_bar_number)],
            ONBOARD_ADDRESS: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_address)],
            ONBOARD_EMERGENCY_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_emergency_name)],
            ONBOARD_EMERGENCY_PHONE: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_emergency_phone)],
            ONBOARD_CONFIRM: [CallbackQueryHandler(confirm_onboarding, pattern='^confirm_onboarding$')],
        },
        fallbacks=[CommandHandler('start', start)],
    )
    
    # Command handlers
    application.add_handler(CommandHandler('start', start))
    application.add_handler(CommandHandler('help', help_command))
    application.add_handler(CommandHandler('quickstart', quickstart))
    application.add_handler(CommandHandler('myagenda', myagenda))
    application.add_handler(CommandHandler('newcase', newcase))
    application.add_handler(CommandHandler('casestatus', casestatus))
    application.add_handler(CommandHandler('requestleave', requestleave))
    application.add_handler(CommandHandler('resources', resources))
    application.add_handler(CommandHandler('emergency', emergency))
    application.add_handler(CommandHandler('refer', refer))
    application.add_handler(CommandHandler('profile', profile))
    application.add_handler(CommandHandler('paymentlink', handle_payment_link_command))

    # Additional callback handlers

    application.add_handler(CallbackQueryHandler(user_guide_callback, pattern='^user_guide$'))
    application.add_handler(CallbackQueryHandler(learn_more_callback, pattern='^learn_more$'))
    application.add_handler(CallbackQueryHandler(quickstart_callback, pattern='^quickstart_cb$'))
    
    application.add_handler(onboarding_handler)

    
    # Web App Data Handler
    application.add_handler(MessageHandler(filters.StatusUpdate.WEB_APP_DATA, handle_web_app_data))
    
    # File Handler
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    
    # Start bot
    logger.info("ðŸš€ City Law Firm Bot is starting...")
    application.run_polling(allowed_updates=Update.ALL_TYPES)




def extract_text_from_file(file_path: str) -> str:
    """Extract text content from various file types"""
    import PyPDF2
    from docx import Document as DocxDocument
    import json
    import markdown
    from openpyxl import load_workbook
    
    file_extension = file_path.lower().split('.')[-1]
    
    try:
        if file_extension == 'pdf':
            # Extract from PDF with improved error handling
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file, strict=False)  # Non-strict mode
                    text = ""
                    num_pages = len(pdf_reader.pages)
                    
                    for i, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                        except Exception as page_error:
                            logger.warning(f"Error extracting page {i+1}: {page_error}")
                            continue
                    
                    if text.strip():
                        return text.strip()
                    else:
                        return f"PDF has {num_pages} pages but no extractable text found. It may be scanned/image-based."
                        
            except Exception as pdf_error:
                logger.error(f"PDF extraction error: {pdf_error}")
                return f"Error reading PDF: {str(pdf_error)}. The file may be corrupted or password-protected."
                
        elif file_extension in ['docx', 'doc']:
            # Extract from Word document
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            if text.strip():
                return text.strip()
            else:
                return "Word document appears to be empty or contains only images."
            
        elif file_extension == 'txt':
            # Extract from text file
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read().strip()
                return content if content else "Text file is empty."
        
        elif file_extension == 'md':
            # Extract from Markdown file
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
                # Convert markdown to plain text (strip formatting)
                html = markdown.markdown(md_content)
                # Simple HTML tag removal
                import re
                text = re.sub('<[^<]+?>', '', html)
                return text.strip() if text.strip() else "Markdown file is empty."
        
        elif file_extension == 'json':
            # Extract from JSON file
            with open(file_path, 'r', encoding='utf-8') as file:
                json_data = json.load(file)
                # Convert JSON to formatted string
                formatted_json = json.dumps(json_data, indent=2)
                return f"JSON Content:\n{formatted_json}"
        
        elif file_extension in ['xlsx', 'xls']:
            # Extract from Excel file
            wb = load_workbook(file_path, data_only=True)
            text = ""
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"\n=== Sheet: {sheet_name} ===\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
            return text.strip() if text.strip() else "Excel file appears to be empty."
                
        else:
            return f"Unsupported file type: {file_extension}. Supported: PDF, DOCX, TXT, MD, JSON, XLSX"
            
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return f"Error extracting text: {str(e)}"



async def analyze_document_with_ai(text_content: str, filename: str) -> str:
    """Analyze document using OpenAI API"""
    from openai import OpenAI
    
    try:
        logger.info(f"Starting AI analysis for {filename}, text length: {len(text_content)} chars")
        
        # Configure OpenAI API
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            return "**AI Analysis Error:**\n\nOpenAI API key not configured. Please add OPENAI_API_KEY to .env file."
        
        client = OpenAI(api_key=api_key)
        
        # Limit text to avoid timeouts (10000 chars ~ 2500 tokens)
        text_sample = text_content[:10000] if len(text_content) > 10000 else text_content
        
        # Create analysis prompt
        prompt = f"""You are a legal document analyzer. Analyze this document and provide:

1. **Document Type**: Identify the type (Contract, Brief, Correspondence, etc.)
2. **Summary**: 2-3 sentence summary
3. **Key Parties**: List important parties/entities
4. **Important Dates**: Extract significant dates
5. **Legal Issues**: Main legal matters
6. **Action Items**: Required actions
7. **Risk Assessment**: Brief risk assessment

Filename: {filename}

Content (first 10000 chars):
{text_sample}

Format your response clearly and concisely."""

        logger.info("Sending request to OpenAI API...")
        
        # Generate analysis with OpenAI
        response = client.chat.completions.create(
            model="gpt-4o-mini",  # Fast and cost-effective
            messages=[
                {"role": "system", "content": "You are a legal document analysis expert."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=1000,
            timeout=60
        )
        
        logger.info("Received response from OpenAI API")
        
        if response and response.choices and response.choices[0].message.content:
            return response.choices[0].message.content
        else:
            return "âš ï¸ **AI Analysis Unavailable**\n\nReceived empty response from AI service."
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Error in AI analysis: {error_msg}")
        return (
            "âš ï¸ **AI Analysis Unavailable**\n\n"
            "The AI service could not be reached to analyze this document. "
            "However, the file has been securely saved to your case files.\n\n"
            f"_Error details: {error_msg}_"
        )


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle file uploads and perform AI analysis"""
    document = update.message.document
    file_id = document.file_id
    file_name = document.file_name
    
    # Download file
    new_file = await context.bot.get_file(file_id)
    file_path = f"downloads/{file_name}"
    os.makedirs("downloads", exist_ok=True)
    await new_file.download_to_drive(file_path)
    
    # Notify user
    await update.message.reply_text(
        f"ðŸ“„ **File Received:** {file_name}\n"
        f"âš–ï¸ **Initial Legal Analysis in progress...**\n\n"
        f"_Extracting text and analyzing content..._",
        parse_mode='Markdown'
    )
    
    # Extract text from file
    text_content = extract_text_from_file(file_path)
    
    if text_content.startswith("Error") or text_content.startswith("Unsupported"):
        await update.message.reply_text(
            f"âŒ {text_content}\n\n"
            f"Supported formats: PDF, DOCX, TXT, MD, JSON, XLSX",
            parse_mode='Markdown'
        )
        return
    
    # Analyze with AI
    ai_summary = await analyze_document_with_ai(text_content, file_name)
    
    # Save to DB
    session = get_session(engine)
    try:
        user = update.effective_user
        db_user = session.query(User).filter_by(telegram_id=user.id).first()
        
        new_doc = Document(
            filename=file_name,
            file_path=file_path,
            file_type=document.mime_type,
            file_size=document.file_size,
            ai_summary=ai_summary,
            uploaded_by=db_user.id if db_user else None
        )
        session.add(new_doc)
        session.commit()
        
        # Store document context for follow-up questions
        context.user_data['last_document'] = {
            'filename': file_name,
            'text_content': text_content,
            'analysis': ai_summary
        }
        
        # Send analysis result with follow-up options
        keyboard = [
            [
                InlineKeyboardButton("ðŸ“¤ Share", switch_inline_query=ai_summary[:50]),
                InlineKeyboardButton("ðŸ’¾ Save to Case", callback_data='doc_save')
            ],
            [InlineKeyboardButton("âœ… Done", callback_data='doc_done')],
            [InlineKeyboardButton("ðŸ’¬ Ask Follow-up Question", callback_data='doc_continue')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            f"âœ… **Initial Legal Analysis Complete**\n\n{ai_summary}",
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        
    except Exception as e:
        logger.error(f"Error saving document: {e}")
        await update.message.reply_text("âŒ Error saving document analysis.")
    finally:
        session.close()


async def handle_payment_link_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/paymentlink [amount] [case_ref] - Generate payment link (Super Admin only)"""
    user = update.effective_user
    super_admin_id = os.getenv('SUPER_ADMIN_ID')
    allowed_username = 'origichidiah'
    
    is_admin = False
    if super_admin_id and str(user.id) == str(super_admin_id):
        is_admin = True
    elif user.username and user.username.lower() == allowed_username.lower():
        is_admin = True

    # Strict Super Admin Check
    if not is_admin:
        await update.message.reply_text(f"â›” **Access Denied**: restricted to Super Admin (Paul Kasim, @{allowed_username}).\nYour ID: `{user.id}`", parse_mode='Markdown')
        return


    # Parse arguments
    if not context.args or len(context.args) < 2:
        await update.message.reply_text(
            "âŒ **Usage:** `/paymentlink [amount] [case_reference]`\n"
            "Example: `/paymentlink 5000 CL-2025-001`",
            parse_mode='Markdown'
        )
        return

    try:
        amount = float(context.args[0])
        case_ref = context.args[1]
        
        # Generate unique token
        token = str(uuid.uuid4())
        
        session = get_session(engine)
        try:
            # Create link record
            db_user = session.query(User).filter_by(telegram_id=user.id).first()
            
            payment_request = PaymentRequest(
                link_token=token,
                amount=amount,
                case_reference=case_ref,
                created_by=db_user.id if db_user else None,
                status='pending'
            )
            
            session.add(payment_request)
            session.commit()
            
            # Construct link - Use API_BASE_URL (ngrok) for external access
            base_url = os.getenv('MINI_APP_URL', 'http://localhost:5000').rsplit('/', 1)[0] # Extract base
            # If MINI_APP_URL is https://.../mini_app/index.html, we want https://...
            if 'mini_app' in base_url:
                 base_url = base_url.split('/mini_app')[0]

            link = f"{base_url}/pay/{token}"
            
            await update.message.reply_text(
                f"âœ… **Payment Link Generated**\n\n"
                f"ðŸ’° **Amount:** ${amount:,.2f}\n"
                f"ðŸ“‚ **Case:** {case_ref}\n"
                f"ðŸ”— **Link:** `{link}`\n\n"
                f"_Send this link to the client securely._",
                parse_mode='Markdown'
            )
            
        except Exception as e:
            await update.message.reply_text(f"âŒ Error DB: {str(e)}")
        finally:
            session.close()

    except ValueError:
             await update.message.reply_text("âŒ Error: Amount must be a number.")
    except Exception as e:
        logger.error(f"Error in payment command: {e}")
        await update.message.reply_text("âŒ An unexpected error occurred.")


async def handle_web_app_data(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle data received from Mini App"""
    data = json.loads(update.effective_message.web_app_data.data)
    action = data.get('action')
    
    if action == 'new_case':
        # Save to database
        session = get_session(engine)
        try:
            user = update.effective_user
            db_user = session.query(User).filter_by(telegram_id=user.id).first()
            
            if not db_user:
                await update.message.reply_text("âŒ Error: Please complete onboarding first with /start")
                return
            
            # Create new case
            new_case = Case(
                case_number=data.get('case_number', f"CL-{datetime.now().strftime('%Y%m%d%H%M')}"),
                title=f"{data.get('client_name', 'New Client')} Case",
                client_name=data.get('client_name', 'Unknown'),
                case_type=data.get('case_type', 'General'),
                status='active',
                priority=data.get('priority', 'normal').lower(),
                assigned_to=db_user.id,
                filing_date=datetime.now(),
                deadline=datetime.now() + timedelta(days=30),
                description=data.get('description', '')
            )
            session.add(new_case)
            session.commit()
            
            await update.message.reply_text(
                f"âœ… **New Case Saved Successfully**\n\n"
                f"ðŸ“‚ Case: {new_case.case_number}\n"
                f"ðŸ‘¤ Client: {new_case.client_name}\n"
                f"âš–ï¸ Type: {new_case.case_type}\n"
                f"ðŸ”¥ Priority: {new_case.priority.title()}\n"
                f"ðŸ‘® Assigned to: {db_user.full_name}\n\n"
                f"Case ID: #{new_case.id}",
                parse_mode='Markdown'
            )
        except Exception as e:
            session.rollback()
            logger.error(f"Error saving new case: {e}")
            await update.message.reply_text(f"âŒ Error saving case: {str(e)}")
        finally:
            session.close()

        
    elif action == 'log_time':
        await update.message.reply_text(
            f"âœ… **Time Entry Logged**\n\n"
            f"â±ï¸ Duration: {data['duration']} hours\n"
            f"ðŸ“ Activity: {data['activity_type']}\n"
            f"ðŸ“„ Description: {data['description']}",
            parse_mode='Markdown'
        )
        
    elif action == 'leave_request':
        # Route to HR/Lead Partner
        await update.message.reply_text(
            f"âœ… **Leave Request Submitted**\n\n"
            f"Your request has been forwarded to **HR** and **Lead Partner** for approval.\n\n"
            f"ðŸ–ï¸ Type: {data['leave_type']}\n"
            f"ðŸ“… From: {data['start_date']}\n"
            f"ðŸ“… To: {data['end_date']}\n"
            f"â“ Reason: {data['reason']}",
            parse_mode='Markdown'
        )
        # e.g., await context.bot.send_message(chat_id=hr_user_id, text=f"New Leave Request from...")

    elif action == 'edit_profile':
        session = get_session(engine)
        try:
            user = update.effective_user
            db_user = session.query(User).filter_by(telegram_id=user.id).first()
            if db_user:
                db_user.full_name = data.get('full_name', db_user.full_name)
                db_user.email = data.get('email', db_user.email)
                db_user.phone = data.get('phone', db_user.phone)
                db_user.address = data.get('address', db_user.address)
                session.commit()
                
                await update.message.reply_text(
                    f"âœ… **Profile Updated Successfully**\n\n"
                    f"ðŸ‘¤ Name: {db_user.full_name}\n"
                    f"ðŸ“§ Email: {db_user.email}\n"
                    f"ðŸ“± Phone: {db_user.phone}\n"
                    f"ðŸ“ Address: {db_user.address}",
                    parse_mode='Markdown'
                )
            else:
                await update.message.reply_text("âŒ Error: User profile not found.")
        except Exception as e:
            logger.error(f"Error updating profile: {e}")
            await update.message.reply_text("âŒ Error updating profile.")
        finally:
            session.close()

    elif action == 'new_agenda_item':
        item_type = data.get('type', 'court')
        if item_type == 'court':
            await update.message.reply_text(
                f"âœ… **Court Date Added**\n\n"
                f"ðŸ“‚ Case: {data.get('case_number', 'N/A')}\n"
                f"ðŸ›ï¸ Court: {data.get('court_name', 'N/A')}\n"
                f"ðŸ“… Date/Time: {data.get('date_time', 'N/A')}\n"
                f"ðŸ“‹ Purpose: {data.get('purpose', 'Hearing')}",
                parse_mode='Markdown'
            )
        else:
            await update.message.reply_text(
                f"âœ… **Task Added**\n\n"
                f"ðŸ“ Title: {data.get('title', 'N/A')}\n"
                f"ðŸ“… Deadline: {data.get('deadline', 'N/A')}\n"
                f"ðŸ”¥ Priority: {data.get('priority', 'medium').capitalize()}",
                parse_mode='Markdown'
            )


async def myagenda_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Wrapper for myagenda callback"""
    query = update.callback_query
    await query.answer()
    
    # Create a synthetic update with message context for myagenda
    # We'll modify the update to have a message attribute pointing to the callback query message
    update._effective_message = query.message
    await myagenda(update, context)

async def help_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Wrapper for help callback"""
    query = update.callback_query
    await query.answer()
    await help_command(update, context)


async def doc_callback_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle document analysis callbacks (Done/Continue)"""
    query = update.callback_query
    await query.answer()
    
    logger.info(f"Document callback received: {query.data}")
    
    if query.data == 'doc_done':
        # Clear document context
        if 'last_document' in context.user_data:
            del context.user_data['last_document']
        if 'awaiting_followup' in context.user_data:
            del context.user_data['awaiting_followup']
            
        await query.edit_message_reply_markup(reply_markup=None)
        await query.message.reply_text("âœ… Document analysis finished. Context cleared.")
        
    elif query.data == 'doc_save':
        # "Save to Case" - For now, just confirm it's in the DB (which happened on upload)
        # In a generic "Save to Case" flow, we might ask which case, but for MVP:
        await query.answer("Document saved to General Files")
        await query.message.reply_text(
            "âœ… **Document Saved**\n"
            "This document has been archived in your general case files.\n"
            "To link it to a specific case, please use the Mini App.",
            parse_mode='Markdown'
        )

    elif query.data == 'doc_continue':
        # Set flag to await follow-up question
        context.user_data['awaiting_followup'] = True
        logger.info(f"Setting awaiting_followup to True for user {update.effective_user.id}")
        
        await query.edit_message_reply_markup(reply_markup=None)
        await query.message.reply_text(
            "ðŸ’¬ **Ask a follow-up question** about the document.\n"
            "I'll use the document context to answer you.",
            parse_mode='Markdown'
        )


async def handle_document_followup(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle text messages for document follow-up questions"""
    # Check if we are awaiting a follow-up and have document context
    is_awaiting = context.user_data.get('awaiting_followup')
    has_context = 'last_document' in context.user_data
    
    logger.info(f"Follow-up handler check - Awaiting: {is_awaiting}, Has Context: {has_context}")
    
    if not is_awaiting or not has_context:
        return # Not a follow-up, let other handlers process it
        
    user_question = update.message.text
    doc_context = context.user_data['last_document']
    
    logger.info(f"Processing follow-up question: {user_question}")
    
    # Notify user
    processing_msg = await update.message.reply_text("ðŸ¤” Analyzing document context...")
    
    try:
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        
        # Create prompt with document context
        prompt = f"""You are a legal assistant helping with a document.
        
Document Filename: {doc_context['filename']}
Document Content (excerpt):
{doc_context['text_content'][:8000]}

Previous Analysis:
{doc_context['analysis']}

User Question: {user_question}

Answer the user's question based on the document content. If the answer is not in the document, state that clearly."""

        # Generate answer
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful legal assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=800
        )
        
        answer = response.choices[0].message.content
        
        # Send answer with options to continue or stop
        keyboard = [
            [InlineKeyboardButton("âœ… Done", callback_data='doc_done')],
            [InlineKeyboardButton("ðŸ’¬ Ask Another Question", callback_data='doc_continue')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await processing_msg.delete()
        await update.message.reply_text(
            f"**Answer:**\n\n{answer}",
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        
        # Reset awaiting flag until they click "Ask Another Question"
        context.user_data['awaiting_followup'] = False
        
    except Exception as e:
        logger.error(f"Error in follow-up: {e}")
        await processing_msg.edit_text("âŒ Error processing your question.")


async def myagenda(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/myagenda - View comprehensive daily schedule and tasks"""
    user = update.effective_user
    session = get_session(engine)
    
    # Handle both command and callback query contexts
    message = update.message if update.message else update.callback_query.message
    
    try:
        db_user = session.query(User).filter_by(telegram_id=user.id).first()
        
        if not db_user or not db_user.onboarding_completed:
            await message.reply_text("âš ï¸ Please complete onboarding first using /start")
            return
        
        today = datetime.now().date()
        week_from_now = today + timedelta(days=7)
        
        # 1. Court Dates (Next 7 days)
        court_dates = session.query(CourtDate).join(Case).filter(
            Case.assigned_to == db_user.id,
            CourtDate.hearing_date >= datetime.now(),
            CourtDate.hearing_date <= datetime.combine(week_from_now, datetime.max.time())
        ).order_by(CourtDate.hearing_date).all()
        
        # 2. Compliance Tasks (Pending)
        tasks = session.query(ComplianceTask).filter(
            ComplianceTask.assigned_to == db_user.id,
            ComplianceTask.status == 'pending'
        ).all()
        
        # 3. Time Entries (Today)
        from sqlalchemy import func
        time_entries = session.query(TimeEntry).filter(
            TimeEntry.user_id == db_user.id,
            func.date(TimeEntry.date) == today
        ).all()
        total_hours = sum(t.duration for t in time_entries)
        
        # Build Agenda Message
        agenda_text = f"ðŸ“… **My Agenda - {today.strftime('%A, %B %d')}**\n\n"
        
        # Court Dates Section
        agenda_text += "âš–ï¸ **Upcoming Court Dates:**\n"
        if court_dates:
            for cd in court_dates:
                date_str = cd.hearing_date.strftime('%b %d %I:%M %p')
                agenda_text += f"â€¢ {date_str}: {cd.case.case_number} - {cd.purpose or 'Hearing'}\n"
        else:
            agenda_text += "_No court dates in the next 7 days._\n"
        
        # Tasks Section
        agenda_text += "\nâœ… **Pending Tasks:**\n"
        if tasks:
            for task in tasks:
                due_str = task.deadline.strftime('%b %d') if task.deadline else "No deadline"
                agenda_text += f"â€¢ {task.title} (Due: {due_str})\n"
        else:
            agenda_text += "_No pending tasks._\n"
            
        # Time Tracking Section
        agenda_text += f"\nâ±ï¸ **Time Logged Today:** {total_hours:.1f} hours\n"
        if time_entries:
            for entry in time_entries:
                agenda_text += f"â€¢ {entry.duration}h - {entry.description}\n"
        
        # Fetch latest broadcast
        latest_broadcast = session.query(Notification).filter_by(notification_type='broadcast').order_by(Notification.created_at.desc()).first()
        broadcast_param = ""
        if latest_broadcast:
            import urllib.parse
            msg = urllib.parse.quote(latest_broadcast.message)
            time = urllib.parse.quote(latest_broadcast.created_at.strftime('%Y-%m-%d %H:%M'))
            broadcast_param = f"?broadcast={msg}&time={time}"

        # Action Buttons
        keyboard = [
            [InlineKeyboardButton("âž• Log Time", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL') + "#time"))],
            [InlineKeyboardButton("ðŸ“‚ New Case", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL') + "#newcase"))],
            [InlineKeyboardButton("ðŸ“Š Dashboard", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL') + broadcast_param))]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await message.reply_text(agenda_text, reply_markup=reply_markup, parse_mode='Markdown')
        
    except Exception as e:
        logger.error(f"Error loading agenda: {e}")
        await message.reply_text("âŒ Error loading agenda.")
    finally:
        session.close()


async def add_to_agenda(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/addagenda - Add new item to your agenda"""
    user = update.effective_user
    
    keyboard = [
        [InlineKeyboardButton("ðŸ“… Court Date", callback_data='add_court_date')],
        [InlineKeyboardButton("âœ… Task/Reminder", callback_data='add_task')],
        [InlineKeyboardButton("â±ï¸ Time Entry", callback_data='add_time_entry')],
        [InlineKeyboardButton("âŒ Cancel", callback_data='cancel_add')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        "ðŸ“ **Add to Agenda**\n\n"
        "What would you like to add?",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def add_court_date_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle adding a court date"""
    query = update.callback_query
    await query.answer()
    
    await query.edit_message_text(
        "ðŸ“… **Add Court Date**\n\n"
        "Please provide the following information in this format:\n\n"
        "`Case Number | Court Name | Date (YYYY-MM-DD HH:MM) | Purpose`\n\n"
        "Example:\n"
        "`CASE-2024-001 | Supreme Court | 2024-12-15 10:00 | Hearing`",
        parse_mode='Markdown'
    )
    
    # Set state to expect court date input
    context.user_data['awaiting_court_date'] = True


async def add_task_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle adding a task"""
    query = update.callback_query
    await query.answer()
    
    await query.edit_message_text(
        "âœ… **Add Task/Reminder**\n\n"
        "Please provide the following information in this format:\n\n"
        "`Task Title | Deadline (DD-MM-YYYY) | Priority (high/medium/low)`\n\n"
        "Example:\n"
        "`Review contract documents | 20-12-2024 | high`",
        parse_mode='Markdown'
    )
    
    # Set state to expect task input
    context.user_data['awaiting_task'] = True


async def process_agenda_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Process text input for adding agenda items"""
    user = update.effective_user
    text = update.message.text
    session = get_session(engine)
    
    try:
        db_user = session.query(User).filter_by(telegram_id=user.id).first()
        
        if not db_user:
            await update.message.reply_text("âš ï¸ Please complete onboarding first using /start")
            return
        
        # Handle court date input
        if context.user_data.get('awaiting_court_date'):
            try:
                parts = [p.strip() for p in text.split('|')]
                if len(parts) != 4:
                    raise ValueError("Invalid format")
                
                case_number, court_name, date_str, purpose = parts
                
                # Find the case
                case = session.query(Case).filter(
                    Case.case_number == case_number,
                    Case.assigned_to == db_user.id
                ).first()
                
                if not case:
                    await update.message.reply_text(
                        f"âŒ Case '{case_number}' not found. Please create the case first."
                    )
                    return
                
                # Parse date
                hearing_date = datetime.strptime(date_str, '%d-%m-%Y %H:%M')
                
                # Create court date
                court_date = CourtDate(
                    case_id=case.id,
                    court_name=court_name,
                    hearing_date=hearing_date,
                    purpose=purpose
                )
                
                session.add(court_date)
                session.commit()
                
                await update.message.reply_text(
                    f"âœ… **Court Date Added!**\n\n"
                    f"ðŸ“… {hearing_date.strftime('%d-%m-%Y %H:%M')}\n"
                    f"ðŸ›ï¸ {court_name}\n"
                    f"ðŸ“‹ Case: {case_number}\n"
                    f"ðŸ“ Purpose: {purpose}",
                    parse_mode='Markdown'
                )
                
                context.user_data['awaiting_court_date'] = False
                
            except ValueError as e:
                await update.message.reply_text(
                    "âŒ Invalid format. Please use:\n"
                    "`Case Number | Court Name | Date (DD-MM-YYYY HH:MM) | Purpose`",
                    parse_mode='Markdown'
                )
            except Exception as e:
                logger.error(f"Error adding court date: {e}")
                await update.message.reply_text("âŒ Error adding court date. Please try again.")
        
        # Handle task input
        elif context.user_data.get('awaiting_task'):
            try:
                parts = [p.strip() for p in text.split('|')]
                if len(parts) != 3:
                    raise ValueError("Invalid format")
                
                title, deadline_str, priority = parts
                
                # Parse deadline
                deadline = datetime.strptime(deadline_str, '%d-%m-%Y')
                
                # Create compliance task
                task = ComplianceTask(
                    title=title,
                    assigned_to=db_user.id,
                    deadline=deadline,
                    priority=priority.lower(),
                    status='pending'
                )
                
                session.add(task)
                session.commit()
                
                priority_emoji = {"high": "ðŸ”´", "medium": "ðŸŸ¡", "low": "ðŸŸ¢"}.get(priority.lower(), "âšª")
                
                await update.message.reply_text(
                    f"âœ… **Task Added!**\n\n"
                    f"ðŸ“‹ {title}\n"
                    f"ðŸ“… Deadline: {deadline.strftime('%d-%m-%Y')}\n"
                    f"{priority_emoji} Priority: {priority.capitalize()}",
                    parse_mode='Markdown'
                )
                
                context.user_data['awaiting_task'] = False
                
            except ValueError as e:
                await update.message.reply_text(
                    "âŒ Invalid format. Please use:\n"
                    "`Task Title | Deadline (DD-MM-YYYY) | Priority (high/medium/low)`",
                    parse_mode='Markdown'
                )
            except Exception as e:
                logger.error(f"Error adding task: {e}")
                await update.message.reply_text("âŒ Error adding task. Please try again.")
    
    finally:
        session.close()


# --- Admin Commands ---

async def promote_admin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Promote self to admin using a secret key"""
    if not context.args:
        await update.message.reply_text("âŒ Usage: /promote_admin <secret_key>")
        return

    secret = context.args[0]
    # In production, use os.getenv('ADMIN_SECRET')
    if secret == "admin123":
        session = get_session(engine)
        try:
            user = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
            if user:
                user.role = 'admin'
                session.commit()
                await update.message.reply_text("âœ… You are now an Admin!")
            else:
                await update.message.reply_text("âŒ User not found. Please start the bot first.")
        finally:
            session.close()
    else:
        await update.message.reply_text("âŒ Invalid secret key.")


async def list_users(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """List all users (Admin only)"""
    session = get_session(engine)
    try:
        # Check admin
        admin = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if not admin or admin.role != 'admin':
            await update.message.reply_text("â›” Admin access required.")
            return

        users = session.query(User).all()
        msg = "ðŸ‘¥ **User List:**\n\n"
        for u in users:
            status_icon = "ðŸŸ¢" if u.status == 'active' else "ðŸ”´"
            msg += f"{status_icon} **{u.full_name}** (ID: `{u.telegram_id}`)\n"
            msg += f"   Role: {u.role} | Status: {u.status}\n\n"
        
        await update.message.reply_text(msg, parse_mode='Markdown')
    finally:
        session.close()


async def block_user(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Block a user (Admin only)"""
    if not context.args:
        await update.message.reply_text("âŒ Usage: /block_user <telegram_id>")
        return
        
    session = get_session(engine)
    try:
        # Check admin
        admin = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if not admin or admin.role != 'admin':
            await update.message.reply_text("â›” Admin access required.")
            return

        target_id = int(context.args[0])
        user = session.query(User).filter_by(telegram_id=target_id).first()
        if user:
            user.status = 'blocked'
            session.commit()
            await update.message.reply_text(f"ðŸš« User {user.full_name} has been BLOCKED.")
        else:
            await update.message.reply_text("âŒ User not found.")
    except ValueError:
        await update.message.reply_text("âŒ Invalid ID format.")
    finally:
        session.close()


async def unblock_user(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Unblock a user (Admin only)"""
    if not context.args:
        await update.message.reply_text("âŒ Usage: /unblock_user <telegram_id>")
        return
        
    session = get_session(engine)
    try:
        # Check admin
        admin = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if not admin or admin.role != 'admin':
            await update.message.reply_text("â›” Admin access required.")
            return

        target_id = int(context.args[0])
        user = session.query(User).filter_by(telegram_id=target_id).first()
        if user:
            user.status = 'active'
            session.commit()
            await update.message.reply_text(f"âœ… User {user.full_name} has been UNBLOCKED.")
        else:
            await update.message.reply_text("âŒ User not found.")
    except ValueError:
        await update.message.reply_text("âŒ Invalid ID format.")
    finally:
        session.close()


async def delete_user(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Delete a user (Admin only)"""
    if not context.args:
        await update.message.reply_text("âŒ Usage: /delete_user <telegram_id>")
        return
        
    session = get_session(engine)
    try:
        # Check admin
        admin = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if not admin or admin.role != 'admin':
            await update.message.reply_text("â›” Admin access required.")
            return

        target_id = int(context.args[0])
        user = session.query(User).filter_by(telegram_id=target_id).first()
        if user:
            # Delete associated records (optional, but good practice)
            # For now just delete user
            session.delete(user)
            session.commit()
            await update.message.reply_text(f"ðŸ—‘ï¸ User {user.full_name} has been DELETED.")
        else:
            await update.message.reply_text("âŒ User not found.")
    except ValueError:
        await update.message.reply_text("âŒ Invalid ID format.")
    finally:
        session.close()

# --- Middleware ---
async def check_block(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Check if user is blocked"""
    if not update.effective_user:
        return
        
    session = get_session(engine)
    try:
        user = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if user and user.status == 'blocked':
            # Stop processing
            raise Application.StopPropagation
    finally:
        session.close()

# --- Profile Actions ---

async def profile(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """View and manage user profile"""
    session = get_session(engine)
    try:
        user = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if not user:
            await update.message.reply_text("âŒ Profile not found. Please /start to onboard.")
            return
            
        msg = (
            f"ðŸ‘¤ **User Profile**\n\n"
            f"**Name:** {user.full_name}\n"
            f"**Role:** {user.role.title()}\n"
            f"**Department:** {user.departments}\n"
            f"**Status:** {user.status.title()}\n\n"
            f"ðŸ“ž **Phone:** {user.phone or 'N/A'}\n"
            f"ðŸ“§ **Email:** {user.email}\n"
            f"ðŸ“ **Address:** {user.address or 'N/A'}\n"
        )
        
        keyboard = [
            [InlineKeyboardButton("âœï¸ Edit Profile", callback_data='edit_profile')],
            [InlineKeyboardButton("ðŸ—‘ï¸ Delete Account", callback_data='delete_account_confirm')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        # Handle both command and callback query
        if update.callback_query:
            await update.callback_query.answer()
            await update.callback_query.edit_message_text(msg, parse_mode='Markdown', reply_markup=reply_markup)
        else:
            await update.message.reply_text(msg, parse_mode='Markdown', reply_markup=reply_markup)
    finally:
        session.close()

async def edit_profile_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle edit profile button"""
    query = update.callback_query
    await query.answer()
    
    # Route based on field selection
    if query.data == 'edit_profile':
        keyboard = [
            [InlineKeyboardButton("ðŸ“± Phone", callback_data='edit_field_phone')],
            [InlineKeyboardButton("ðŸ“§ Email", callback_data='edit_field_email')],
            [InlineKeyboardButton("ðŸ“ Address", callback_data='edit_field_address')],
            [InlineKeyboardButton("ðŸ”™ Cancel", callback_data='profile_back')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text("Select a field to edit:", reply_markup=reply_markup)
    elif query.data == 'edit_field_phone':
        await query.edit_message_text("ðŸ“± Please enter your new phone number:")
        return EDIT_PHONE
    elif query.data == 'edit_field_email':
        await query.edit_message_text("ðŸ“§ Please enter your new email address:")
        return EDIT_EMAIL
    elif query.data == 'edit_field_address':
        await query.edit_message_text("ðŸ“ Please enter your new address:")
        return EDIT_ADDRESS
    
    return ConversationHandler.END

async def handle_edit_phone(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle phone number update"""
    new_phone = update.message.text.strip()
    session = get_session(engine)
    try:
        user = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if user:
            user.phone = new_phone
            session.commit()
            await update.message.reply_text(f"âœ… Phone number updated to: {new_phone}")
        else:
            await update.message.reply_text("âŒ User not found.")
    finally:
        session.close()
    return ConversationHandler.END

async def handle_edit_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle email update"""
    new_email = update.message.text.strip()
    
    # Basic email validation
    if '@' not in new_email or '.' not in new_email:
        await update.message.reply_text("âš ï¸ Please enter a valid email address:")
        return EDIT_EMAIL
    
    session = get_session(engine)
    try:
        user = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if user:
            user.email = new_email
            session.commit()
            await update.message.reply_text(f"âœ… Email updated to: {new_email}")
        else:
            await update.message.reply_text("âŒ User not found.")
    finally:
        session.close()
    return ConversationHandler.END

async def handle_edit_address(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle address update"""
    new_address = update.message.text.strip()
    session = get_session(engine)
    try:
        user = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if user:
            user.address = new_address
            session.commit()
            await update.message.reply_text(f"âœ… Address updated to: {new_address}")
        else:
            await update.message.reply_text("âŒ User not found.")
    finally:
        session.close()
    return ConversationHandler.END

async def delete_account_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle delete account confirmation"""
    query = update.callback_query
    await query.answer()
    
    if query.data == 'delete_account_confirm':
        keyboard = [
            [InlineKeyboardButton("âœ… Yes, Delete My Account", callback_data='delete_account_final')],
            [InlineKeyboardButton("âŒ No, Cancel", callback_data='profile_back')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            "âš ï¸ **Are you sure?**\n\n"
            "This will permanently delete your account and all your data. This action cannot be undone.",
            parse_mode='Markdown',
            reply_markup=reply_markup
        )
    elif query.data == 'delete_account_final':
        session = get_session(engine)
        try:
            user = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
            if user:
                session.delete(user)
                session.commit()
                await query.edit_message_text("ðŸ—‘ï¸ **Account Deleted.**\n\nGoodbye!")
            else:
                await query.edit_message_text("âŒ Error: User not found.")
        finally:
            session.close()
    elif query.data == 'profile_back':
        # Redirect to profile view
        await query.delete_message()
        await profile(update, context)

# --- Time Logging ---

async def logtime(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Log billable hours: /logtime <hours> <case_number> <description>"""
    if not context.args or len(context.args) < 3:
        await update.message.reply_text(
            "âŒ **Usage:** `/logtime <hours> <case_number> <description>`\n"
            "Example: `/logtime 1.5 CL-2025-089 Research on precedent`",
            parse_mode='Markdown'
        )
        return

    try:
        hours = float(context.args[0])
        case_number = context.args[1]
        description = " ".join(context.args[2:])
    except ValueError:
        await update.message.reply_text("âŒ Invalid hours. Please use a number (e.g., 1.5).")
        return

    session = get_session(engine)
    try:
        user = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if not user:
            await update.message.reply_text("âŒ User not found.")
            return

        # Find case
        case = session.query(Case).filter_by(case_number=case_number).first()
        if not case:
            await update.message.reply_text(f"âŒ Case `{case_number}` not found.", parse_mode='Markdown')
            return

        # Create Time Entry
        entry = TimeEntry(
            user_id=user.id,
            case_id=case.id,
            date=datetime.utcnow(),
            duration_minutes=int(hours * 60),
            hourly_rate=250.0, # Default rate
            activity_type="General",
            description=description,
            billable=True
        )
        session.add(entry)
        session.commit()
        
        await update.message.reply_text(
            f"âœ… **Time Logged!**\n\n"
            f"â±ï¸ {hours} hours ({int(hours*60)} mins)\n"
            f"ðŸ“‚ Case: {case.title}\n"
            f"ðŸ“ {description}",
            parse_mode='Markdown'
        )
        
    finally:
        session.close()

# --- Notification Board ---

async def broadcast(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Send a broadcast notification to all users (Admin only)"""
    if not context.args:
        await update.message.reply_text("âŒ Usage: /broadcast <message>")
        return
        
    message_text = " ".join(context.args)
    
    session = get_session(engine)
    try:
        # Check admin
        admin = session.query(User).filter_by(telegram_id=update.effective_user.id).first()
        if not admin or admin.role != 'admin':
            await update.message.reply_text("â›” Admin access required.")
            return

        # Create Notification Record
        notification = Notification(
            title="ðŸ“¢ System Broadcast",
            message=message_text,
            notification_type="broadcast",
            created_by=admin.id,
            sent=True,
            sent_at=datetime.utcnow()
        )
        session.add(notification)
        session.commit()
        
        # Send to all users
        users = session.query(User).filter(User.status == 'active').all()
        count = 0
        for user in users:
            try:
                await context.bot.send_message(
                    chat_id=user.telegram_id,
                    text=f"ðŸ“¢ **ANNOUNCEMENT**\n\n{message_text}",
                    parse_mode='Markdown'
                )
                count += 1
            except Exception:
                pass 
        
        await update.message.reply_text(f"âœ… Broadcast sent to {count} users.")
        
    finally:
        session.close()

async def setup_commands(application: Application):
    """Set up bot commands"""
    from telegram import BotCommand
    commands = [
        BotCommand("start", "Start the bot"),
        BotCommand("help", "Show help message"),
        BotCommand("myagenda", "View your agenda"),
        BotCommand("addagenda", "Add item to agenda"),
        BotCommand("newcase", "Create a new case"),
        BotCommand("casestatus", "Check case status"),
        BotCommand("logtime", "Log billable hours"),
        BotCommand("profile", "Manage profile"),
        BotCommand("resources", "Access legal resources"),
        BotCommand("emergency", "Emergency contacts"),
        BotCommand("refer", "Refer a client"),
        BotCommand("broadcast", "ðŸ“¢ Send broadcast (Admin)"),
        BotCommand("list_users", "ðŸ‘¥ List users (Admin)"),
    ]
    await application.bot.set_my_commands(commands)


async def start_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle back to start callback"""
    query = update.callback_query
    await query.answer()
    
    user = update.effective_user
    session = get_session(engine)
    try:
        db_user = session.query(User).filter_by(telegram_id=user.id).first()
        if not db_user:
             await query.edit_message_text("Please run /start to onboard.")
             return

        # Determine greeting
        greeting = "ðŸ‘‹ Welcome back"
        
        broadcast_param = ""
        # Check for unread notifications (simplified check)
        # ...

        keyboard = [
            [InlineKeyboardButton("ðŸ“± Open Virtual Office", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL') + broadcast_param))],
            [InlineKeyboardButton("ðŸ“‹ My Agenda", callback_data='my_agenda')],
            [InlineKeyboardButton("ðŸ“Š Dashboard", callback_data='dashboard')],
            [InlineKeyboardButton("â„¹ï¸ Help", callback_data='help')]
        ]
        
        await query.edit_message_text(
            f"{greeting}, {db_user.full_name}!\n\n"
            f"ðŸ¢ **City Law Firm Virtual Office**\n"
            f"Departments: {db_user.departments}\n"
            f"Position: {db_user.position}\n\n"
            f"What would you like to do today?",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
    finally:
        session.close()

async def dashboard_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle dashboard callback"""
    query = update.callback_query
    await query.answer()
    
    user = update.effective_user
    session = get_session(engine)
    try:
        db_user = session.query(User).filter_by(telegram_id=user.id).first()
        today = datetime.now().date()
        week_from_now = datetime.now() + timedelta(days=7)
        
        # Count active cases
        active_cases = session.query(Case).filter(
            Case.assigned_to == db_user.id,
            Case.status.in_(['open', 'in_progress', 'active'])
        ).count()
        
        # Upcoming court dates
        upcoming_court = session.query(CourtDate).join(Case).filter(
            Case.assigned_to == db_user.id,
            CourtDate.hearing_date >= datetime.now(),
            CourtDate.hearing_date <= week_from_now
        ).count()
        
        # Pending leave
        pending_leave = session.query(LeaveRequest).filter(
            LeaveRequest.user_id == db_user.id,
            LeaveRequest.status == 'pending'
        ).count()
        
        dashboard_text = (
            f"ðŸ“Š **Dashboard - {db_user.full_name}**\n\n"
            f"ðŸ¢ **Dept:** {db_user.departments}\n"
            f"ðŸ’¼ **Role:** {db_user.position}\n\n"
            f"**ðŸ“ˆ Your Stats:**\n"
            f"â€¢ Active Cases: {active_cases}\n"
            f"â€¢ Court Dates (7d): {upcoming_court}\n"
            f"â€¢ Pending Leave: {pending_leave}\n\n"
            f"_Last updated: {datetime.now().strftime('%H:%M')}_"
        )
        
        keyboard = [
            [InlineKeyboardButton("ðŸ“± Open Virtual Office", web_app=WebAppInfo(url=os.getenv('MINI_APP_URL')))],
            [InlineKeyboardButton("ðŸ“‹ My Agenda", callback_data='my_agenda')],
            [InlineKeyboardButton("ðŸ”„ Refresh", callback_data='dashboard')],
            [InlineKeyboardButton("Â« Back", callback_data='back_to_start')]
        ]
        
        await query.edit_message_text(
            dashboard_text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
    except Exception as e:
        logger.error(f"Dashboard error: {e}")
        await query.edit_message_text("âŒ Error loading dashboard.")
    finally:
        session.close()

def main():
    """Start the bot"""
    # Create application
    persistence = PicklePersistence(filepath='conversationbot')
    application = Application.builder().token(os.getenv('BOT_TOKEN')).persistence(persistence).post_init(setup_commands).build()
    
    # Conversation handler
    onboarding_handler = ConversationHandler(
        entry_points=[CallbackQueryHandler(start_onboarding_callback, pattern='^start_onboarding$')],
        states={
            ONBOARD_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_name)],
            ONBOARD_EMAIL: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_email)],
            ONBOARD_PHONE: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_phone)],
            ONBOARD_DEPARTMENT: [CallbackQueryHandler(onboard_department, pattern='^dept_')],
            ONBOARD_POSITION: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_position)],
            ONBOARD_SPECIALIZATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_specialization)],
            ONBOARD_BAR_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_bar_number)],
            ONBOARD_ADDRESS: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_address)],
            ONBOARD_EMERGENCY_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_emergency_name)],
            ONBOARD_EMERGENCY_PHONE: [MessageHandler(filters.TEXT & ~filters.COMMAND, onboard_emergency_phone)],
            ONBOARD_PHOTO: [MessageHandler(filters.PHOTO, onboard_photo)],
            ONBOARD_LOCATION: [MessageHandler(filters.LOCATION, onboard_location)],
            ONBOARD_CONFIRM: [CallbackQueryHandler(confirm_onboarding, pattern='^confirm_onboarding$')],
        },
        fallbacks=[CommandHandler('start', start)],
        name="onboarding",
        persistent=True,
    )
    
    # Middleware (Check blocked status first)
    from telegram.ext import TypeHandler
    application.add_handler(TypeHandler(Update, check_block), group=-1)

    # Callback Handlers
    application.add_handler(CallbackQueryHandler(start_callback, pattern='^back_to_start$'))
    application.add_handler(CallbackQueryHandler(dashboard_callback, pattern='^dashboard$'))
    application.add_handler(CallbackQueryHandler(myagenda_callback, pattern='^my_agenda$'))
    application.add_handler(CallbackQueryHandler(help_callback, pattern='^help$'))

    # Command handlers
    application.add_handler(CommandHandler('start', start))
    
    # Admin Commands
    application.add_handler(CommandHandler('promote_admin', promote_admin))
    application.add_handler(CommandHandler('list_users', list_users))
    application.add_handler(CommandHandler('block_user', block_user))
    application.add_handler(CommandHandler('unblock_user', unblock_user))
    application.add_handler(CommandHandler('delete_user', delete_user))
    application.add_handler(CommandHandler('logtime', logtime))
    application.add_handler(CommandHandler('broadcast', broadcast))
    
    application.add_handler(CommandHandler('help', help_command))
    application.add_handler(CommandHandler('myagenda', myagenda))
    application.add_handler(CommandHandler('addagenda', add_to_agenda))
    application.add_handler(CommandHandler('newcase', newcase))
    application.add_handler(CommandHandler('casestatus', casestatus))
    application.add_handler(CommandHandler('requestleave', requestleave))
    application.add_handler(CommandHandler('resources', resources))
    application.add_handler(CommandHandler('emergency', emergency))
    application.add_handler(CommandHandler('refer', refer))
    application.add_handler(CommandHandler('profile', profile))
    
    # Profile editing conversation handler
    profile_edit_handler = ConversationHandler(
        entry_points=[CallbackQueryHandler(edit_profile_callback, pattern='^edit_field_')],
        states={
            EDIT_PHONE: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_edit_phone)],
            EDIT_EMAIL: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_edit_email)],
            EDIT_ADDRESS: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_edit_address)],
        },
        fallbacks=[CallbackQueryHandler(profile, pattern='^profile_back$')],
        name="profile_edit",
        persistent=True,
    )
    application.add_handler(profile_edit_handler)
    
    # Profile Callbacks
    application.add_handler(CallbackQueryHandler(edit_profile_callback, pattern='^edit_profile$'))
    application.add_handler(CallbackQueryHandler(delete_account_callback, pattern='^delete_account_'))
    application.add_handler(CallbackQueryHandler(profile, pattern='^profile_back$'))
    
    # File Handler (Register BEFORE ConversationHandler to ensure it catches files)
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    
    # Agenda management callbacks (before onboarding)
    application.add_handler(CallbackQueryHandler(add_court_date_callback, pattern='^add_court_date$'))
    application.add_handler(CallbackQueryHandler(add_task_callback, pattern='^add_task$'))

    # Web App Data Handler
    application.add_handler(MessageHandler(filters.StatusUpdate.WEB_APP_DATA, handle_web_app_data))
    
    # Document Analysis Callbacks
    application.add_handler(CallbackQueryHandler(doc_callback_handler, pattern='^doc_'))

    # Onboarding Handler (Must be BEFORE generic MessageHandlers)
    application.add_handler(onboarding_handler)
    
    # Document interaction callbacks
    application.add_handler(CallbackQueryHandler(doc_callback_handler, pattern='^doc_'))
    
    # Document follow-up handler (must be before other text handlers)
    application.add_handler(MessageHandler(
        filters.TEXT & ~filters.COMMAND & ~filters.Regex('^/'), 
        handle_document_followup
    ))
    
    # Agenda input handler (must be after other text handlers but before onboarding)
    # NOTE: Changed to be AFTER onboarding to avoid capturing onboarding inputs
    application.add_handler(MessageHandler(
        filters.TEXT & ~filters.COMMAND & ~filters.Regex('^/'),
        process_agenda_input
    ))
    
    # Web App Data Handler
    application.add_handler(MessageHandler(filters.StatusUpdate.WEB_APP_DATA, handle_web_app_data))
    
    # Start bot
    logger.info("ðŸš€ City Law Firm Bot is starting...")
    application.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == '__main__':
    main()

